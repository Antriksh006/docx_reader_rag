from docx import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from groq import Groq
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)


def create_vector_store(docx_file_path, embedding_function):
    """Create a vector store from a DOCX file."""
    text = extract_text_from_docx(docx_file_path)

    # Split the text into chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=500)
    documents = text_splitter.split_text(text)

    # Create Chroma vector store
    db = Chroma.from_texts(documents, embedding_function)

    return db


def get_bot_response(user_input, db, embedding_function, groq_client):
    """Get the bot response based on user input."""
    # Generate embedding for the user query
    query_embedding = embedding_function.embed_query(user_input)

    # Perform similarity search using Chroma
    similar_chunks = db.similarity_search_by_vector(query_embedding)

    # Gather context from similar chunks
    context = " ".join(chunk.page_content for chunk in similar_chunks)

    # Construct the detailed prompt
    detailed_prompt = f"You are a question-answering chatbot. Answer the following question: {user_input} \nContext: {context}"

    # Make a call to Groq API for chat completions
    chat_completion = groq_client.chat.completions.create(
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": detailed_prompt}
        ],
        model="llama3-8b-8192",
        max_tokens=1000
    )

    return chat_completion.choices[0].message.content


# Function to split text into lines by "\n"
def split_text_by_lines(text_list):
    """Splits each element in the tth list by '\n' and returns a list of lines."""
    lines = []
    for text in text_list:
        lines.extend(text.split('\n'))
    return lines


# Updated highlight function that works line by line
def highlight_line_by_line(paragraph, line_to_highlight):
    """
    Function to highlight a single line in a paragraph.
    """
    # Normalize the paragraph text by replacing line breaks and trimming spaces
    paragraph_text = paragraph.text.replace("\n", " ").strip()

    # Clean up the line to avoid issues with special characters
    cleaned_line = line_to_highlight.replace("\n", " ").strip()

    # Check if the cleaned line exists in the paragraph text
    if cleaned_line in paragraph_text:
        remaining_text = cleaned_line  # Track remaining text to be highlighted

        # Loop through the runs to find where the line is located
        for run in paragraph.runs:
            run_text = run.text.replace("\n", " ").strip()

            if remaining_text and remaining_text.startswith(run_text):
                # Highlight the part of the run that matches
                highlight_element = OxmlElement('w:highlight')
                highlight_element.set(qn('w:val'), 'yellow')  # Highlight color
                run._r.get_or_add_rPr().append(highlight_element)

                # Update the remaining text to check for the rest of the line
                remaining_text = remaining_text[len(run_text):]
            else:
                remaining_text = cleaned_line if cleaned_line in run_text else None


def get_text_to_highlight(user_input, db, embedding_function):
    """Get the text to highlight based on user input."""
    # Generate embedding for the user query
    query_embedding = embedding_function.embed_query(user_input)

    # Perform similarity search using Chroma
    tth = []
    similar_chunks = db.similarity_search_by_vector(query_embedding)

    for chunk in similar_chunks:
        texttohighlight = chunk.page_content
        tth.append(texttohighlight)
    return tth


def main():
    """Main function to run the chatbot and highlight relevant text."""
    # Get user inputs
    user_input = input("ENTER_YOUR_QUERY_HERE: ")
    API_KEY = input("ENTER_YOUR_GROQ_API_KEY_HERE: ")
    docx_file_path = input("ENTER_FILE_PATH: ")

    # Initialize Groq client
    groq_client = Groq(api_key=API_KEY)

    # Initialize embedding function
    embedding_function = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

    # Create the vector store
    db = create_vector_store(docx_file_path, embedding_function)

    # Get and print the bot response
    response = get_bot_response(user_input, db, embedding_function, groq_client)
    print(response)

    do_you_want_citation = input("Do you want citation for this text? (Enter 1 for confirmation): ")

    if do_you_want_citation == "1":
        doc = Document(docx_file_path)

        # Example tth array with text to highlight
        tth = get_text_to_highlight(user_input, db, embedding_function)

        # Split tth elements into lines
        tth_lines = split_text_by_lines(tth)

        # Loop through each paragraph and try to highlight line by line
        for paragraph in doc.paragraphs:
            for line in tth_lines:
                highlight_line_by_line(paragraph, line)

        # Save the modified document
        doc.save('highlighted_file_by_lines.docx')
        os.startfile('highlighted_file_by_lines.docx')


if __name__ == "__main__":
    main()
